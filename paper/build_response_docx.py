import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex="F0F0F0"):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def create_response_document(output_path):
    doc = docx.Document()
    
    # Margins
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("Response to Reviewers' Comments")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0, 51, 102)

    # Subtitle / Metadata
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("7th International Conference on Electronics and Sustainable Communication Systems (ICESC 2026)")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True

    # Paper Details Box
    t_meta = doc.add_table(rows=4, cols=2)
    t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_info = [
        ("Paper Title:", "Multi-Agent LLM Orchestration for Claim Authentication via Semantic Codebase Vectorization"),
        ("Authors:", "M. Mounika, B. Lavanya, Abhinav Basam, G. Uday Kiran, B. Priyanka, B. VeeraSekhar"),
        ("Affiliation:", "Dept. of Computer Science & Engineering (AI & ML), B V Raju Institute of Technology, India"),
        ("Decision:", "Accepted with Minor Revisions / Camera-Ready Submission")
    ]
    for r_idx, (k, v) in enumerate(meta_info):
        c0 = t_meta.cell(r_idx, 0)
        c1 = t_meta.cell(r_idx, 1)
        set_cell_shading(c0, "EAECEE")
        set_cell_margins(c0, 40, 40, 60, 60)
        set_cell_margins(c1, 40, 40, 60, 60)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.name = "Arial"
        r0.font.size = Pt(9.5)
        r0.font.bold = True
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.name = "Arial"
        r1.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Intro letter
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(8)
    p_intro.paragraph_format.line_spacing = 1.15
    r_intro = p_intro.add_run(
        "We thank the Technical Program Committee and the Reviewers for their constructive evaluation and insightful feedback. "
        "We have carefully revised the manuscript to adhere to all IEEE camera-ready formatting guidelines and to address the technical review comments. "
        "A detailed, point-by-point response to the reviewers' comments is provided below."
    )
    r_intro.font.name = "Arial"
    r_intro.font.size = Pt(10)

    def add_sec_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title)
        r.font.name = "Arial"
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 51, 102)

    def add_comment_block(comm_title, comment_text, response_text, changes_text):
        # Comment Box
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_before = Pt(8)
        p_c.paragraph_format.space_after = Pt(2)
        r_ct = p_c.add_run(comm_title + "\n")
        r_ct.font.name = "Arial"
        r_ct.font.size = Pt(10.5)
        r_ct.font.bold = True
        r_ct.font.color.rgb = RGBColor(150, 0, 0)
        
        r_c = p_c.add_run(f'"{comment_text}"')
        r_c.font.name = "Arial"
        r_c.font.size = Pt(9.5)
        r_c.font.italic = True

        # Response
        p_r = doc.add_paragraph()
        p_r.paragraph_format.space_before = Pt(2)
        p_r.paragraph_format.space_after = Pt(4)
        p_r.paragraph_format.line_spacing = 1.15
        
        r_rt = p_r.add_run("Author Response: ")
        r_rt.font.name = "Arial"
        r_rt.font.size = Pt(10)
        r_rt.font.bold = True
        r_rt.font.color.rgb = RGBColor(0, 102, 51)
        
        r_resp = p_r.add_run(response_text)
        r_resp.font.name = "Arial"
        r_resp.font.size = Pt(9.5)

        # Changes in Manuscript
        p_ch = doc.add_paragraph()
        p_ch.paragraph_format.space_before = Pt(2)
        p_ch.paragraph_format.space_after = Pt(8)
        
        r_cht = p_ch.add_run("Modifications in Manuscript: ")
        r_cht.font.name = "Arial"
        r_cht.font.size = Pt(9.5)
        r_cht.font.bold = True
        
        r_ch = p_ch.add_run(changes_text)
        r_ch.font.name = "Arial"
        r_ch.font.size = Pt(9.5)

    # SECTION 1: REVIEW COMMENTS 1
    add_sec_heading("I. Responses to Reviewer 1 Comments")

    add_comment_block(
        "Comment 1.1: Benchmark Evaluation Dataset Size",
        "The proposed CodeAudit AI framework is evaluated using only 30 benchmark test cases, and validation on larger real-world recruitment datasets with diverse candidate profiles would better demonstrate its robustness and scalability.",
        "We agree with the reviewer's suggestion to validate the framework on a larger, more diverse dataset. In response, we have incorporated an Extended Validation Cohort comprising 150 real-world candidate profiles. These profiles were drawn from open-source contributors with varying portfolio sizes (ranging from 2 to 45 repositories and 12 to over 1,200 source files per candidate).",
        "Added Section IV-A(2) and Table IV to present the evaluation results on the extended cohort, comparing the proposed framework against baseline models."
    )

    add_comment_block(
        "Comment 1.2: Comparison with SOTA Code Models & RAG Frameworks",
        "Although the framework is compared with a keyword-matching baseline, it is not evaluated against recent state-of-the-art code understanding models, modern multi-agent systems, or advanced RAG-based resume screening frameworks.",
        "To provide a thorough comparative analysis, we have benchmarked the proposed framework against recent state-of-the-art AI resume screening frameworks and advanced code understanding foundation models, including CodeBERT, UniXcoder, Code Llama, and StarCoder.",
        "Added Table I (Section II) and Table XI (Section IV-J) to summarize the comparative analysis against existing screening approaches."
    )

    add_comment_block(
        "Comment 1.3: Practical Computational Profiling & Latency",
        "The practical performance of the framework, such as inference time, response latency, computational cost per candidate audit, and memory footprint during vector indexing, is not thoroughly analyzed.",
        "We have conducted a detailed computational profiling of the framework to assess its practical viability. The profiling measures inference time, memory footprint, and estimated cost per candidate audit across a representative workload.",
        "Added Section IV-H and Table IX, which detail the end-to-end latency (7.27 s), peak memory usage (187 MB), and estimated API cost per audit."
    )

    add_comment_block(
        "Comment 1.4: Private Repositories, Limited Open Source & Fairness",
        "The paper does not adequately discuss how the framework handles candidates with private repositories, limited open-source activity, or proprietary enterprise codebases, which may impact real-world fairness.",
        "We acknowledge the importance of addressing private repositories and fairness in recruitment. The revised manuscript outlines how the system accommodates private enterprise codebases via scoped OAuth authorization and sanitized local uploads. Furthermore, we clarify that unverified claims are classified as 'Unsubstantiated' rather than explicitly fraudulent, providing a basis for subsequent human verification without penalizing the candidate.",
        "Expanded Section V-B (Fairness and Private Repository Compliance) and Section VI (Conclusion and Future Work)."
    )

    add_comment_block(
        "Comment 1.5: Multi-Agent Pipeline Component Ablation & LLM Backbones",
        "The contribution of individual components within the multi-agent pipeline is not isolated through ablation studies, making it unclear whether performance gains arise from RAG indexing, the LLM Judge, or specific prompt rubrics.",
        "To isolate the contributions of individual pipeline components, we have included an ablation study evaluating the performance of each processing phase. Additionally, we have incorporated an evaluation of various LLM backbones to determine their respective impacts on the final verification accuracy.",
        "Added Table III (Multi-Agent Pipeline Component Ablation) and Table V (Comparative Evaluation of LLM Judge Model Backbones) in Sections IV-B and IV-D, respectively."
    )

    # SECTION 2: REVIEW COMMENTS 2
    add_sec_heading("II. Responses to Reviewer 2 Comments")

    add_comment_block(
        "Comment 2.1: RAG Chunk Size & Overlap Windowing Rationale",
        "Provide clear technical justification and empirical sensitivity analysis for the selected 800-character chunk window and 100-character sliding stride.",
        "We have provided empirical justification for the selected RAG parameters through a sensitivity analysis of chunk window sizes and overlap strides. The analysis demonstrates that an 800-character window provides an optimal balance between context preservation and vector density.",
        "Added Section IV-F and Table VII detailing the chunk size sensitivity evaluation."
    )

    add_comment_block(
        "Comment 2.2: Prompt Sensitivity & Temperature Robustness",
        "Evaluate the framework's sensitivity to prompt variations, reasoning instructions, and LLM temperature hyperparameters.",
        "We have introduced an evaluation of prompt sensitivity and temperature robustness. The analysis compares different reasoning strategies (Zero-Shot, Few-Shot, and Chain-of-Thought) and sampling temperatures to assess their impact on deterministic reproducibility and precision.",
        "Added Section IV-G and Table VIII to present the results of the prompt sensitivity and temperature robustness analysis."
    )

    add_comment_block(
        "Comment 2.3: Integration with ATS Systems & Webhook Architecture",
        "Clarify how the framework integrates with enterprise Applicant Tracking Systems (ATS) and human-in-the-loop workflows.",
        "We have elaborated on the framework's integration with enterprise Applicant Tracking Systems (ATS). The revised architecture specifies the mechanisms for asynchronous processing, webhook dispatching, and direct candidate record updating.",
        "Expanded Section III-F to outline the four-level ATS integration architecture."
    )

    add_comment_block(
        "Comment 2.4: Embedding Model Rationale (all-MiniLM-L6-v2)",
        "Justify why all-MiniLM-L6-v2 was selected over larger specialized code embedding backbones.",
        "The selection of the all-MiniLM-L6-v2 embedding model is now explicitly justified. The rationale highlights its advantages in low vector dimensionality, computational efficiency on standard CPUs, and cross-domain transfer capabilities for natural-language-to-code retrieval.",
        "Added detailed technical justification in Section III-D."
    )

    add_comment_block(
        "Comment 2.5: Dense Embedding Model Comparison",
        "Include a comparative evaluation of multiple embedding model backbones on retrieval accuracy and latency.",
        "To support our model selection, we have conducted a comparative evaluation of several dense embedding models, analyzing their retrieval accuracy (Hit@5, MRR), CPU latency, and memory footprint.",
        "Added Section IV-E and Table VI, which summarize the empirical comparison of the dense embedding models."
    )

    add_comment_block(
        "Comment 2.6: Academic Tone & Citation Integrity",
        "Ensure professional academic style, remove banned colloquial phrases ('holistic', 'you'), and ensure all references are peer-reviewed and cited in-text.",
        "The manuscript has been rigorously reviewed to ensure a professional academic tone. We have eliminated colloquial phrasing and ensured that all citations adhere strictly to the IEEE formatting and numbering standards.",
        "Performed a comprehensive textual revision and verified the formatting of the 30 included references."
    )

    try:
        doc.save(output_path)
        print(f"Successfully generated Response to Reviewers document at: {output_path}")
    except Exception as e:
        print(f"Could not save {output_path}: {e} (Close Word to overwrite)")

if __name__ == "__main__":
    create_response_document(r"C:\Users\abhin_yr5b44s\Desktop\resume_project\paper\Response_to_Review_Comments_ICESC2026.docx")
    create_response_document(r"C:\Users\abhin_yr5b44s\Downloads\Response_to_Review_Comments_ICESC2026.docx")
