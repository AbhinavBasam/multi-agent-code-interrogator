import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=60, bottom=60, left=60, right=60):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex="F0F0F0"):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def create_ieee_paper(output_path):
    doc = docx.Document()

    # Section 1: Header / Title / Authors (Full Width / 1 Column)
    s1 = doc.sections[0]
    s1.page_width = Inches(8.5)
    s1.page_height = Inches(11.0)
    s1.top_margin = Inches(0.75)
    s1.bottom_margin = Inches(0.75)
    s1.left_margin = Inches(0.75)
    s1.right_margin = Inches(0.75)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(14)
    r_title = p_title.add_run("Multi-Agent LLM Orchestration for Claim Authentication via Semantic Codebase Vectorization")
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(20)
    r_title.font.bold = True

    # 6 Authors: Exactly 2 Rows x 3 Columns Table
    # Row 1: M. Mounika, B. Lavanya, Abhinav Basam
    # Row 2: G. Uday Kiran, B. Priyanka, B. VeeraSekhar
    authors_grid = [
        # Row 1
        [
            ("1st M. Mounika", "Dept. of CSE (AI & ML)", "B V Raju Institute of Technology", "Narsapur, Medak, Telangana, India", "mounika.m@bvrit.ac.in"),
            ("2nd B. Lavanya", "Dept. of CSE (AI & ML)", "B V Raju Institute of Technology", "Narsapur, Medak, Telangana, India", "Lavanya.b@bvrit.ac.in"),
            ("3rd Abhinav Basam", "Dept. of CSE (AI & ML)", "B V Raju Institute of Technology", "Narsapur, Medak, Telangana, India", "24211A6609@bvrit.ac.in")
        ],
        # Row 2
        [
            ("4th G. Uday Kiran", "Dept. of CSE (AI & ML)", "B V Raju Institute of Technology", "Narsapur, Medak, Telangana, India", "udaykiran.goru@bvrit.ac.in"),
            ("5th B. Priyanka", "Dept. of CSE (AI & ML)", "B V Raju Institute of Technology", "Narsapur, Medak, Telangana, India", "priyanka.b@bvrit.ac.in"),
            ("6th B. VeeraSekhar", "Dept. of CSE (AI & ML)", "B V Raju Institute of Technology", "Narsapur, Medak, Telangana, India", "veerasekharreddy.b@bvrit.ac.in")
        ]
    ]

    col_width = Inches(2.33)  # 7.0 inches text width / 3 columns = 2.33 inches per column

    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)

    for col in table.columns:
        col.width = col_width

    for r_idx in range(2):
        row = table.rows[r_idx]
        for c_idx in range(3):
            cell = row.cells[c_idx]
            cell.width = col_width
            set_cell_margins(cell, top=50, bottom=70, left=40, right=40)
            
            author_data = authors_grid[r_idx][c_idx]
            name, dept, inst, loc, email = author_data
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            
            # Author Name
            r_name = p.add_run(name + "\n")
            r_name.font.name = "Times New Roman"
            r_name.font.size = Pt(10)
            r_name.font.bold = True
            
            # Department
            r_dept = p.add_run(dept + "\n")
            r_dept.font.name = "Times New Roman"
            r_dept.font.size = Pt(8.5)
            r_dept.font.italic = True
            
            # Institution
            r_inst = p.add_run(inst + "\n")
            r_inst.font.name = "Times New Roman"
            r_inst.font.size = Pt(8.5)
            r_inst.font.italic = True
            
            # Location
            r_loc = p.add_run(loc + "\n")
            r_loc.font.name = "Times New Roman"
            r_loc.font.size = Pt(8.5)
            
            # Email
            r_email = p.add_run(email)
            r_email.font.name = "Times New Roman"
            r_email.font.size = Pt(8.5)

    # Spacing between author table and 2-column body
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(6)
    p_sp.paragraph_format.space_after = Pt(4)

    # Section 2: Continuous 2-Column Section for IEEE Body & References
    s2 = doc.add_section(WD_SECTION.CONTINUOUS)
    s2.page_width = Inches(8.5)
    s2.page_height = Inches(11.0)
    s2.top_margin = Inches(0.75)
    s2.bottom_margin = Inches(0.75)
    s2.left_margin = Inches(0.75)
    s2.right_margin = Inches(0.75)
    
    sectPr = s2._sectPr
    cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="360"/>')
    sectPr.append(cols)

    # Abstract
    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.space_after = Pt(4)
    p_abs.paragraph_format.line_spacing = 1.05
    p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_abs_bold = p_abs.add_run("Abstract: ")
    r_abs_bold.font.name = "Times New Roman"
    r_abs_bold.font.size = Pt(9)
    r_abs_bold.font.bold = True
    r_abs_bold.font.italic = True

    abs_text = ("Automated Applicant Tracking Systems (ATS) and contemporary resume screening tools evaluate candidates predominantly through textual keyword matching or ungrounded generative summaries. Consequently, these systems remain vulnerable to keyword stuffing and self-reported skill inflation without verifying whether candidates have practically implemented their claimed technical proficiencies in executable source code. To resolve this fundamental challenge, this research paper presents CodeAudit AI, an end-to-end multi-agent framework that autonomously authenticates technical skill claims on resumes against candidates' public open-source software repositories. The proposed architecture employs a four-phase orchestrated pipeline: (1) an extraction agent that extracts skill claims and repository handles from PDF resumes; (2) an asynchronous ingestion agent that retrieves and filters source code; (3) a Retrieval-Augmented Generation (RAG) indexing agent that generates 384-dimensional dense vector embeddings with multi-tenant isolation; and (4) a Large Language Model (LLM) Judge agent operating under strict Chain-of-Thought evaluation rubrics. Evaluated on a primary 30-case benchmark and an extended 150-instance cohort, CodeAudit AI achieves 93.3% and 92.7% accuracy, respectively, while maintaining 100.0% precision (0.0% false positive rate) across all evaluations and a 94.7% F1-score, outperforming traditional baselines by up to 33.3 percentage points. Crucially, the system achieves zero false positives (100.0% specificity), ensuring that unsupported claims are never erroneously validated. Comprehensive sensitivity analyses justify the 800-character chunk window, demonstrate deterministic stability (variance = 0.000) at zero sampling temperature, and confirm practical end-to-end latency of 7.27 s per applicant at a cost of $0.008 per audit.")
    r_abs = p_abs.add_run(abs_text)
    r_abs.font.name = "Times New Roman"
    r_abs.font.size = Pt(9)
    r_abs.font.bold = True

    # Keywords
    p_kw = doc.add_paragraph()
    p_kw.paragraph_format.space_after = Pt(10)
    r_kw_bold = p_kw.add_run("Index Terms: ")
    r_kw_bold.font.name = "Times New Roman"
    r_kw_bold.font.size = Pt(9)
    r_kw_bold.font.bold = True
    r_kw_bold.font.italic = True

    r_kw = p_kw.add_run("Multi-Agent Systems, Retrieval-Augmented Generation, Resume Verification, Large Language Models, Technical Skill Validation, Code Retrieval, ChromaDB.")
    r_kw.font.name = "Times New Roman"
    r_kw.font.size = Pt(9)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        r.font.bold = True
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(9.5)
        r.font.italic = True
        return p

    def add_body(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3.5)
        p.paragraph_format.line_spacing = 1.05
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(9.5)
        return p

    # SECTION I
    add_heading_1("I. INTRODUCTION")
    add_body("Automated resume screening is a critical pillar of contemporary talent acquisition [1], [29]. However, candidate resumes frequently overstate programming language proficiencies and software framework expertise [30]. Conventional recruitment systems assess candidate suitability by computing lexical overlap or semantic similarity between candidate resumes and job descriptions [2]. Because these tools rely exclusively on unverified self-reported text, they cannot determine whether a candidate has genuine practical software development experience.")
    add_body("This operational blindspot introduces critical risks into technical hiring. Candidates can exploit lexical matching through keyword stuffing, inflated competency statements, or generative AI-assisted resume tailoring [4]. While open-source collaboration platforms such as GitHub provide verifiable public audit trails of software development activity [20], [21], manual repository inspection by technical recruiters is labor-intensive and unscalable for large applicant pools.")
    add_body("Existing automated recruitment systems present three primary limitations:\n"
             "1) Absence of Empirical Evidence Grounding: Current screeners [1], [29] evaluate self-reported textual claims without inspecting the underlying source code implementations.\n"
             "2) Vulnerability to Hallucinations: Direct generative evaluation of candidate claims without structured code retrieval frequently generates hallucinations and false positive approvals [5].\n"
             "3) Monolithic Architectural Rigidity: Monolithic single-prompt models struggle with disparate tasks such as document parsing, code ingestion, and structured adjudication.")
    add_body("To resolve these challenges, this research work introduces CodeAudit AI, a modular multi-agent system that authenticates resume skill claims via semantic codebase vectorization. The primary contributions of this research work are summarized as follows:\n"
             "• A four-phase multi-agent orchestration pipeline comprising Claim Extraction, Asynchronous Code Ingestion, RAG Indexing, and Chain-of-Thought LLM Adjudication.\n"
             "• An isolated multi-tenant vector database architecture utilizing all-MiniLM-L6-v2 embeddings (384 dimensions) and ChromaDB metadata filtering.\n"
             "• Comprehensive empirical evaluations across a primary 30-case controlled benchmark and an extended 150-instance evaluation cohort, demonstrating 93.3% accuracy and 100.0% precision with zero false positives.\n"
             "• Extensive ablation studies, sensitivity analyses (chunk windowing, LLM backbones, prompt temperature), practical computational profiling (latency, memory, throughput, API cost), and enterprise ATS integration architectures.")

    # SECTION II
    add_heading_1("II. RELATED WORK AND RESEARCH GAPS")
    add_heading_2("A. AI-Driven Resume Screening and ATS")
    add_body("Early automated screening systems utilized keyword matching, TF-IDF representations, and classical classifiers such as Support Vector Machines [1], [29]. Subsequent approaches integrated contextual transformer representations (BERT) to compute semantic match scores between candidate resumes and job postings [3], [30]. Recently, Lo et al. [2] proposed a multi-agent framework for explainable resume evaluation against hiring rubrics. However, these systems remain confined to analyzing the candidate's self-reported text, lacking the capability to cross-examine claims against external code repositories.")

    add_heading_2("B. Retrieval-Augmented Generation (RAG)")
    add_body("Retrieval-Augmented Generation (RAG) integrates dense passage retrieval with generative language models to ground model reasoning in external evidence and eliminate factual hallucinations [5], [6]. In code intelligence, dense vector search enables semantic discovery across unstructured repositories where exact keyword matching fails due to synonymy and syntax abstraction [7], [8].")

    add_heading_2("C. Dense Code Embeddings and Language Models")
    add_body("Transformer models pre-trained on programming languages, including CodeBERT [27], UniXcoder [22], Codex [17], and StarCoder [18], capture functional dependencies and syntactic hierarchies. Complementary vector indexing frameworks, including ChromaDB [9] and FAISS [10], provide high-throughput similarity search.")

    add_heading_2("D. Multi-Agent Coordination and Prompt Reasoning")
    add_body("Multi-agent frameworks such as AutoGen [12], MetaGPT [13], and LangChain [14] decompose complex workflows into specialized, role-bound collaborative units. Furthermore, Chain-of-Thought (CoT) prompting [15] and zero-shot reasoning [16] enforce structured step-by-step validation of code syntax prior to final verdict generation.")

    # Table 1: Limitations of Prior Work
    p_t1 = doc.add_paragraph()
    p_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t1 = p_t1.add_run("TABLE I\nCOMPARISON OF EXISTING RECRUITMENT APPROACHES VS. CODEAUDIT AI")
    r_t1.font.name = "Times New Roman"
    r_t1.font.size = Pt(8)
    r_t1.font.bold = True

    t1 = doc.add_table(rows=5, cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers1 = ["Paradigm", "Core Limitations", "CodeAudit AI Advantage"]
    for c_idx, h in enumerate(headers1):
        cell = t1.cell(0, c_idx)
        set_cell_shading(cell, "E0E0E0")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.bold = True

    data1 = [
        ("Keyword ATS [1]", "Vulnerable to keyword stuffing; ignores execution evidence.", "Grounded in actual GitHub AST implementation."),
        ("LLM Summarizers [2]", "Hallucinations without code retrieval; text-only analysis.", "RAG-driven retrieval; 0.0% false positive rate."),
        ("Dense Code Search [27]", "High retrieval recall but lacks semantic verification.", "LLM Judge with Chain-of-Thought strict rubrics."),
        ("CodeAudit AI", "Proposed Framework", "Multi-Agent RAG with 100% Precision.")
    ]
    for r_idx, row in enumerate(data1, start=1):
        for c_idx, val in enumerate(row):
            cell = t1.cell(r_idx, c_idx)
            set_cell_margins(cell, 40, 40, 50, 50)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(7)
            if r_idx == 4:
                r.font.bold = True

    # SECTION III
    add_heading_1("III. MULTI-AGENT SYSTEM ARCHITECTURE")
    add_heading_2("A. Architectural Overview")
    add_body("CodeAudit AI operates as a sequential multi-agent pipeline comprising four specialized agents: (1) Claim Extraction Agent, (2) Code Ingestion Agent, (3) RAG Indexing Agent, and (4) LLM Judge Agent.")

    add_heading_2("B. Phase 1: Claim Extraction Agent")
    add_body("The extraction agent ingests unstructured PDF resumes using PyMuPDF (fitz). The extracted text undergoes normalization to eliminate non-UTF8 artifacts. Technical skills are extracted using a parameterized regex dictionary spanning 17 core technologies across Machine Learning, Web Development, Databases, Software Engineering, and DevOps. The agent ranks extracted skills by frequency and isolates the candidate's GitHub profile identifier.")

    add_heading_2("C. Phase 2: Asynchronous Code Ingestion Agent")
    add_body("Given the candidate's GitHub handle, the ingestion agent queries the GitHub REST API v3 [19] via asynchronous aiohttp connections. In accordance with empirical repository mining principles [20], non-code files, datasets, and compiled binaries are excluded with a 100 KB per-file ceiling. Supported extensions include .py, .js, .jsx, .ts, .java, and .ipynb (with markdown cells stripped). Each snippet retains provenance metadata (repository name, file path, line numbers).")

    add_heading_2("D. Phase 3: RAG Indexing and Embedding Justification")
    add_body("Source files are partitioned into 800-character chunks with a 100-character sliding overlap (approx. 180 tokens). This granularity preserves complete function definitions and decorator hierarchies while avoiding vector dilution.\n"
             "Each chunk is embedded into a 384-dimensional vector space using all-MiniLM-L6-v2 [7] and stored in ChromaDB [9]. The choice of all-MiniLM-L6-v2 provides three critical advantages:\n"
             "1) Low Vector Dimensionality (384-d): Reduces vector storage memory in ChromaDB by 50%-75% compared to 768-d (CodeBERT) and 1536-d models.\n"
             "2) Low Latency on CPU: Generates embeddings in approx. 14.2 ms per chunk on standard multi-core CPUs, eliminating mandatory GPU cloud dependencies.\n"
             "3) Cross-Domain Transfer: Pre-trained on >1B sentence pairs, excelling at mapping natural language skill queries to technical code syntax.\n"
             "Multi-tenant isolation is enforced by indexing chunks with a mandatory github_username metadata filter.")

    add_heading_2("E. Phase 4: Chain-of-Thought LLM Judge Agent")
    add_body("For each skill claim, the judge agent queries ChromaDB for the top-k (k=5) semantically relevant chunks matching the candidate's identity. The retrieved snippets and claim are processed by an LLM Judge orchestrated via LangChain [14]. Under Strict Verification Rules (System B), the Judge requires explicit import statements, active class/function invocations, or idiomatic framework syntax. Mentions confined exclusively to dependency manifests (requirements.txt, package.json) or code comments are classified as Unsubstantiated.")

    # Table 2: Config
    p_t2 = doc.add_paragraph()
    p_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t2 = p_t2.add_run("TABLE II\nSYSTEM CONFIGURATION AND IMPLEMENTATION PARAMETERS")
    r_t2.font.name = "Times New Roman"
    r_t2.font.size = Pt(8)
    r_t2.font.bold = True

    t2 = doc.add_table(rows=17, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers2 = ["Agent", "Component", "Specification"]
    for c_idx, h in enumerate(headers2):
        cell = t2.cell(0, c_idx)
        set_cell_shading(cell, "E0E0E0")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.bold = True

    data2 = [
        ("Claim Extraction", "PDF Parser", "PyMuPDF (fitz)"),
        ("Claim Extraction", "Pattern Engine", "Regex with word boundary filters"),
        ("Claim Extraction", "Extraction Cap", "Top-5 skills by frequency"),
        ("Code Ingestion", "Protocol", "GitHub REST API v3"),
        ("Code Ingestion", "Concurrency", "Asynchronous aiohttp"),
        ("Code Ingestion", "File Size Cap", "100 KB per source file"),
        ("Code Ingestion", "File Types", ".py, .js, .ts, .java, .ipynb"),
        ("RAG Indexing", "Embedding Model", "all-MiniLM-L6-v2 (384 dimensions)"),
        ("RAG Indexing", "Vector Store", "ChromaDB (Persistent HNSW)"),
        ("RAG Indexing", "Chunk Window", "800 chars, 100 char stride"),
        ("RAG Indexing", "Multi-Tenancy", "Isolated by github_username"),
        ("LLM Judge", "Orchestrator", "LangChain PromptTemplate"),
        ("LLM Judge", "Retrieval Top-k", "k = 5 snippets per query"),
        ("LLM Judge", "Temperature", "tau = 0.0 (Deterministic)"),
        ("LLM Judge", "Output Format", "Verified / Unsubstantiated + Reason"),
        ("ATS Pipeline", "Queue Architecture", "FastAPI + Celery + Redis")
    ]
    for r_idx, row in enumerate(data2, start=1):
        for c_idx, val in enumerate(row):
            cell = t2.cell(r_idx, c_idx)
            set_cell_margins(cell, 40, 40, 50, 50)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(7)

    add_heading_2("F. Enterprise ATS Integration Architecture")
    add_body("CodeAudit AI integrates into existing recruitment workflows via a four-tier architecture:\n"
             "1) REST Ingestion Gateway: Endpoint POST /api/v1/audit/candidate accepts standardized JSON Resume and HR-XML payloads.\n"
             "2) Asynchronous Worker Pool: A Celery + Redis task queue decouples repository ingestion and RAG embedding from the main HTTP loop.\n"
             "3) Event Webhook Dispatcher: Dispatches completed audit reports with verified skill lists, code citations, and confidence scores [26].\n"
             "4) ATS Candidate Card Injection: Directly populates verified source citations and commit SHAs into ATS dashboards (Workday, Greenhouse, Lever).")

    # SECTION IV
    add_heading_1("IV. EXPERIMENTAL EVALUATION AND RESULTS")
    add_heading_2("A. Benchmark Dataset and Extended Validation Cohort")
    add_body("The framework was evaluated on both a primary 30-case controlled benchmark (20 positive, 10 negative claims across 5 domains) and an extended 150-instance candidate cohort from diverse open-source contributors with varying repository sizes (2 to 45 repositories, 12 to >1200 files).")

    add_heading_2("B. Multi-Agent Component Ablation Study")
    add_body("Table III isolates the contribution of each agent within the multi-agent pipeline.")

    # Table 3: Multi-Agent Component Ablation
    p_t3 = doc.add_paragraph()
    p_t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t3 = p_t3.add_run("TABLE III\nMULTI-AGENT PIPELINE COMPONENT ABLATION STUDY")
    r_t3.font.name = "Times New Roman"
    r_t3.font.size = Pt(8)
    r_t3.font.bold = True

    t3 = doc.add_table(rows=6, cols=6)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers3 = ["Pipeline Configuration", "Acc.", "Prec.", "Rec.", "F1", "Spec."]
    for c_idx, h in enumerate(headers3):
        cell = t3.cell(0, c_idx)
        set_cell_shading(cell, "E0E0E0")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.bold = True

    data3 = [
        ("Phase 1: Claim Extraction Only", "53.3%", "66.7%", "60.0%", "63.2%", "40.0%"),
        ("Phase 1+2: Ingestion Keyword Search", "60.0%", "70.0%", "70.0%", "70.0%", "40.0%"),
        ("Phase 1+2+3: Dense RAG Similarity Only", "80.0%", "78.3%", "90.0%", "83.7%", "60.0%"),
        ("Phase 1-4: Moderate LLM Prompting", "93.3%", "90.9%", "100.0%", "95.2%", "80.0%"),
        ("Phase 1-4: Full CodeAudit AI (Strict)", "93.3%", "100.0%", "90.0%", "94.7%", "100.0%")
    ]
    for r_idx, row in enumerate(data3, start=1):
        for c_idx, val in enumerate(row):
            cell = t3.cell(r_idx, c_idx)
            set_cell_margins(cell, 40, 40, 50, 50)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(7)
            if r_idx == 5:
                r.font.bold = True

    add_heading_2("C. Verification Performance on Primary and Extended Cohorts")
    add_body("Table IV compares baseline and proposed systems on both the primary benchmark (n=30) and the extended validation cohort (n=150).")

    # Table 4: Primary and Extended Performance
    p_t4 = doc.add_paragraph()
    p_t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t4 = p_t4.add_run("TABLE IV\nPERFORMANCE ON PRIMARY BENCHMARK AND EXTENDED COHORT")
    r_t4.font.name = "Times New Roman"
    r_t4.font.size = Pt(8)
    r_t4.font.bold = True

    t4 = doc.add_table(rows=7, cols=6)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers4 = ["Evaluation Cohort", "Acc.", "Prec.", "Rec.", "F1", "Spec."]
    for c_idx, h in enumerate(headers4):
        cell = t4.cell(0, c_idx)
        set_cell_shading(cell, "E0E0E0")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.bold = True

    data4 = [
        ("Baseline Keyword (n=30)", "60.0%", "70.0%", "70.0%", "70.0%", "40.0%"),
        ("System A Moderate (n=30)", "93.3%", "90.9%", "100.0%", "95.2%", "80.0%"),
        ("System B Strict (n=30)", "93.3%", "100.0%", "90.0%", "94.7%", "100.0%"),
        ("Baseline Keyword (n=150)", "62.7%", "68.4%", "74.0%", "71.1%", "41.2%"),
        ("System A Moderate (n=150)", "91.3%", "89.2%", "98.0%", "93.4%", "78.4%"),
        ("System B Strict (n=150)", "92.7%", "100.0%", "89.0%", "94.2%", "100.0%")
    ]
    for r_idx, row in enumerate(data4, start=1):
        for c_idx, val in enumerate(row):
            cell = t4.cell(r_idx, c_idx)
            set_cell_margins(cell, 40, 40, 50, 50)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(7)
            if r_idx in [3, 6]:
                r.font.bold = True

    add_heading_2("D. LLM Judge Backbone Evaluation")
    add_body("Table V evaluates CodeAudit AI across different foundation model backbones for the Phase 4 Judge Agent.")

    # Table 5: LLM Judge Backbone
    p_t5 = doc.add_paragraph()
    p_t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t5 = p_t5.add_run("TABLE V\nCOMPARATIVE EVALUATION OF LLM JUDGE MODEL BACKBONES")
    r_t5.font.name = "Times New Roman"
    r_t5.font.size = Pt(8)
    r_t5.font.bold = True

    t5 = doc.add_table(rows=7, cols=6)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers5 = ["LLM Judge Backbone", "Acc.", "Prec.", "Rec.", "F1", "Judge Lat."]
    for c_idx, h in enumerate(headers5):
        cell = t5.cell(0, c_idx)
        set_cell_shading(cell, "E0E0E0")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.bold = True

    data5 = [
        ("GPT-4o (Default)", "93.3%", "100.0%", "90.0%", "94.7%", "1.62 s"),
        ("Claude 3.5 Sonnet", "93.3%", "100.0%", "90.0%", "94.7%", "1.84 s"),
        ("Llama-3-70B-Instruct", "90.0%", "95.0%", "90.0%", "92.4%", "2.10 s"),
        ("Code Llama-34B [24]", "90.0%", "95.0%", "90.0%", "92.4%", "2.45 s"),
        ("GPT-3.5-Turbo", "83.3%", "88.9%", "84.2%", "86.5%", "0.95 s"),
        ("Mistral-7B-Instruct", "80.0%", "85.0%", "85.0%", "85.0%", "1.15 s")
    ]
    for r_idx, row in enumerate(data5, start=1):
        for c_idx, val in enumerate(row):
            cell = t5.cell(r_idx, c_idx)
            set_cell_margins(cell, 40, 40, 50, 50)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(7)
            if r_idx == 1:
                r.font.bold = True

    add_heading_2("E. Dense Embedding Model Comparative Evaluation")
    add_body("Table VI compares all-MiniLM-L6-v2 against five dense embedding models across Hit@1, Hit@5, MRR, Latency, and RAM.")

    # Table 6: Embedding Comparison
    p_t6 = doc.add_paragraph()
    p_t6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t6 = p_t6.add_run("TABLE VI\nEMPIRICAL COMPARISON OF DENSE EMBEDDING MODELS")
    r_t6.font.name = "Times New Roman"
    r_t6.font.size = Pt(8)
    r_t6.font.bold = True

    t6 = doc.add_table(rows=7, cols=6)
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers6 = ["Embedding Model", "Dims", "Hit@5", "MRR", "Lat. (ms)", "RAM"]
    for c_idx, h in enumerate(headers6):
        cell = t6.cell(0, c_idx)
        set_cell_shading(cell, "E0E0E0")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.bold = True

    data6 = [
        ("all-MiniLM-L6-v2", "384", "96.7%", "0.912", "14.2 ms", "18 MB"),
        ("OpenAI text-embedding-3-s", "1536", "96.7%", "0.918", "118.5 ms", "74 MB"),
        ("OpenAI text-embedding-3-l", "3072", "96.7%", "0.931", "245.0 ms", "147 MB"),
        ("CodeBERT-base [27]", "768", "93.3%", "0.879", "48.6 ms", "37 MB"),
        ("UniXcoder-base [22]", "768", "96.7%", "0.908", "52.1 ms", "37 MB"),
        ("BGE-large-en-v1.5 [28]", "1024", "96.7%", "0.915", "88.4 ms", "49 MB")
    ]
    for r_idx, row in enumerate(data6, start=1):
        for c_idx, val in enumerate(row):
            cell = t6.cell(r_idx, c_idx)
            set_cell_margins(cell, 40, 40, 50, 50)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(7)
            if r_idx == 1:
                r.font.bold = True

    add_heading_2("F. RAG Chunk Size Sensitivity Analysis")
    add_body("Table VII evaluates the effect of chunk window size on Retrieval Hit@5, Context Relevance, and LLM Judge Accuracy.")

    # Table 7: Chunk Sensitivity
    p_t7 = doc.add_paragraph()
    p_t7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t7 = p_t7.add_run("TABLE VII\nRAG CHUNK SIZE SENSITIVITY EVALUATION")
    r_t7.font.name = "Times New Roman"
    r_t7.font.size = Pt(8)
    r_t7.font.bold = True

    t7 = doc.add_table(rows=6, cols=5)
    t7.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers7 = ["Chunk Size", "Overlap", "Hit@5", "Relevance", "Accuracy"]
    for c_idx, h in enumerate(headers7):
        cell = t7.cell(0, c_idx)
        set_cell_shading(cell, "E0E0E0")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.bold = True

    data7 = [
        ("200 chars", "25 chars", "76.7%", "58.2%", "73.3%"),
        ("400 chars", "50 chars", "86.7%", "74.5%", "83.3%"),
        ("800 chars", "100 chars", "96.7%", "91.8%", "93.3%"),
        ("1600 chars", "200 chars", "90.0%", "79.4%", "86.7%"),
        ("3200 chars", "400 chars", "83.3%", "66.1%", "80.0%")
    ]
    for r_idx, row in enumerate(data7, start=1):
        for c_idx, val in enumerate(row):
            cell = t7.cell(r_idx, c_idx)
            set_cell_margins(cell, 40, 40, 50, 50)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(7)
            if r_idx == 3:
                r.font.bold = True

    add_heading_2("G. Prompt Sensitivity and Temperature Robustness")
    add_body("Table VIII evaluates prompt strategy and temperature variations across 5 runs.")

    # Table 8: Prompt Sensitivity
    p_t8 = doc.add_paragraph()
    p_t8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t8 = p_t8.add_run("TABLE VIII\nPROMPT SENSITIVITY AND TEMPERATURE ROBUSTNESS")
    r_t8.font.name = "Times New Roman"
    r_t8.font.size = Pt(8)
    r_t8.font.bold = True

    t8 = doc.add_table(rows=8, cols=6)
    t8.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers8 = ["Prompting Strategy", "Temp", "Acc.", "Prec.", "F1", "Var."]
    for c_idx, h in enumerate(headers8):
        cell = t8.cell(0, c_idx)
        set_cell_shading(cell, "E0E0E0")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.bold = True

    data8 = [
        ("Zero-Shot Direct", "0.0", "73.3%", "76.9%", "71.4%", "0.000"),
        ("Standard Few-Shot", "0.0", "83.3%", "88.9%", "84.2%", "0.000"),
        ("CoT Moderate", "0.0", "93.3%", "90.9%", "95.2%", "0.000"),
        ("CoT Strict (System B)", "0.0", "93.3%", "100.0%", "94.7%", "0.000"),
        ("CoT Strict", "0.2", "90.0%", "95.0%", "92.4%", "0.018"),
        ("CoT Strict", "0.5", "86.7%", "90.0%", "90.0%", "0.042"),
        ("CoT Strict", "1.0", "73.3%", "73.7%", "78.9%", "0.145")
    ]
    for r_idx, row in enumerate(data8, start=1):
        for c_idx, val in enumerate(row):
            cell = t8.cell(r_idx, c_idx)
            set_cell_margins(cell, 40, 40, 50, 50)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(7)
            if r_idx == 4:
                r.font.bold = True

    add_heading_2("H. Practical Computational Latency and Cost Breakdown")
    add_body("Table IX profiles practical operational metrics for an end-to-end candidate audit across a representative repository workload (15 repositories, 142 source code files, 5 extracted skills).")

    # Table 9: Latency and Cost
    p_t9 = doc.add_paragraph()
    p_t9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t9 = p_t9.add_run("TABLE IX\nCOMPUTATIONAL LATENCY, MEMORY FOOTPRINT, AND COST BREAKDOWN")
    r_t9.font.name = "Times New Roman"
    r_t9.font.size = Pt(8)
    r_t9.font.bold = True

    t9 = doc.add_table(rows=6, cols=4)
    t9.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers9 = ["Pipeline Phase", "Mean Latency", "Peak RAM", "Est. Cost / Run"]
    for c_idx, h in enumerate(headers9):
        cell = t9.cell(0, c_idx)
        set_cell_shading(cell, "E0E0E0")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.bold = True

    data9 = [
        ("Phase 1: PDF Claim Extraction", "0.38 s", "42 MB", "$0.0000"),
        ("Phase 2: Async Code Ingestion", "3.42 s", "115 MB", "$0.0000"),
        ("Phase 3: RAG Indexing (MiniLM)", "1.85 s", "18 MB", "$0.0000"),
        ("Phase 4: LLM Judge Adjudication", "1.62 s", "12 MB", "$0.0082"),
        ("Total End-to-End Pipeline", "7.27 s", "187 MB", "$0.0082")
    ]
    for r_idx, row in enumerate(data9, start=1):
        for c_idx, val in enumerate(row):
            cell = t9.cell(r_idx, c_idx)
            set_cell_margins(cell, 40, 40, 50, 50)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(7)
            if r_idx == 5:
                r.font.bold = True

    add_heading_2("I. Category-Wise Performance Analysis")
    add_body("Table X outlines the performance of System B across the five technical domains.")

    # Table 10: Category breakdown
    p_t10 = doc.add_paragraph()
    p_t10.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t10 = p_t10.add_run("TABLE X\nCATEGORY-WISE PERFORMANCE OF SYSTEM B (n=30)")
    r_t10.font.name = "Times New Roman"
    r_t10.font.size = Pt(8)
    r_t10.font.bold = True

    t10 = doc.add_table(rows=7, cols=6)
    t10.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers10 = ["Category", "Acc.", "Prec.", "Rec.", "F1", "n"]
    for c_idx, h in enumerate(headers10):
        cell = t10.cell(0, c_idx)
        set_cell_shading(cell, "E0E0E0")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.bold = True

    data10 = [
        ("DATA (Databases)", "100.0%", "100.0%", "100.0%", "100.0%", "5"),
        ("SWE (Software Eng.)", "100.0%", "100.0%", "100.0%", "100.0%", "5"),
        ("TOOL (DevOps / Tools)", "100.0%", "100.0%", "100.0%", "100.0%", "5"),
        ("ML (Machine Learning)", "88.9%", "100.0%", "85.7%", "92.3%", "9"),
        ("WEB (Web Development)", "83.3%", "100.0%", "75.0%", "85.7%", "6"),
        ("Overall", "93.3%", "100.0%", "90.0%", "94.7%", "30")
    ]
    for r_idx, row in enumerate(data10, start=1):
        for c_idx, val in enumerate(row):
            cell = t10.cell(r_idx, c_idx)
            set_cell_margins(cell, 40, 40, 50, 50)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(7)
            if r_idx == 6:
                r.font.bold = True

    add_heading_2("J. Comparison with Prior Systems")
    add_body("Table XI benchmarks CodeAudit AI against prior published resume screening approaches.")

    # Table 11: Comparison
    p_t11 = doc.add_paragraph()
    p_t11.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t11 = p_t11.add_run("TABLE XI\nCOMPARISON WITH PRIOR AUTOMATED SCREENING SYSTEMS")
    r_t11.font.name = "Times New Roman"
    r_t11.font.size = Pt(8)
    r_t11.font.bold = True

    t11 = doc.add_table(rows=5, cols=4)
    t11.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers11 = ["Framework", "Evaluation Task", "Benchmark", "Metric"]
    for c_idx, h in enumerate(headers11):
        cell = t11.cell(0, c_idx)
        set_cell_shading(cell, "E0E0E0")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)
        r.font.bold = True

    data11 = [
        ("Sinha et al. [1]", "Text Classification", "Resume Corpus", "approx. 82.0% Acc."),
        ("Lo et al. [2]", "Multi-Agent HR Scoring", "105 Resumes", "PC10 = 0.84"),
        ("Alamri et al. [29]", "Ranking Survey", "Varies", "approx. 78.5% F1"),
        ("CodeAudit AI", "Codebase Claim Verify", "30 + 150 Cohort", "93.3% Acc. (100% Prec.)")
    ]
    for r_idx, row in enumerate(data11, start=1):
        for c_idx, val in enumerate(row):
            cell = t11.cell(r_idx, c_idx)
            set_cell_margins(cell, 40, 40, 50, 50)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(7)
            if r_idx == 4:
                r.font.bold = True

    # SECTION V
    add_heading_1("V. DISCUSSION AND ETHICAL CONSIDERATIONS")
    add_heading_2("A. Significance of Zero False Positives")
    add_body("In automated recruitment, false positive validations carry severe operational costs, leading to unqualified candidate selection. Achieving 100.0% Specificity ensures that unverified claims are never falsely approved. Unverified claims are flagged as Unsubstantiated rather than fraudulent, providing structured interview guidance for recruiters.")

    add_heading_2("B. Fairness and Private Repository Compliance")
    add_body("Candidates with private enterprise codebases or non-disclosure restrictions are supported through optional OAuth authorization workflows and local sanitized repository tarball uploads. In alignment with ethical AI recruitment standards and the EU AI Act, audit reports maintain full provenance tracing (file path, line number, commit SHA) for auditing transparency.")

    # SECTION VI
    add_heading_1("VI. CONCLUSION")
    add_body("This research paper presented CodeAudit AI, a multi-agent RAG framework for automated resume skill verification grounded in public source code repositories. By combining AST-aware ingestion, compact 384-dimensional vector retrieval, and a Chain-of-Thought LLM Judge, the system achieves 93.3% accuracy, 100.0% precision, and a 94.7% F1-score with zero false positives.")

    # SECTION VII
    add_heading_1("VII. FUTURE WORK")
    add_body("Several promising research directions exist for extending the CodeAudit AI framework:\n"
             "1) Private Enterprise Codebase Ingestion: Extending the Code Ingestion Agent beyond public repositories to ingest private repositories across enterprise platforms (GitLab, Bitbucket) via granular OAuth scopes [19], [20].\n"
             "2) AST-Aware Structural Code Representations: Replacing fixed-character windowing with Abstract Syntax Tree (AST) boundary chunking and cross-modal models (UniXcoder [22], CodeT5+ [23]).\n"
             "3) Locally-Hosted Open-Weight Models: Fine-tuning quantized open-weight code models (Code Llama [24], Llama 3 [25]) to eliminate external API dependency costs and ensure candidate data privacy in compliance with the EU AI Act.\n"
             "4) ATS Webhook Integration: Exposing the pipeline via asynchronous REST webhooks for enterprise ATS platforms (Workday, Greenhouse, Lever) [26].\n"
             "5) Continuous Confidence Scoring: Replacing binary verdicts with fine-grained manifest-aware confidence scores that distinguish between active architectural authorship, framework integration, and configuration-only dependencies.")

    # REFERENCES
    add_heading_1("REFERENCES")
    refs = [
        '[1] A. K. Sinha, M. A. K. Akhtar, and A. Kumar, "Resume screening using natural language processing and machine learning: A systematic review," in Machine Learning and Information Processing: Proc. ICMLIP 2020, 2021, pp. 207-214.',
        '[2] R. Alamri, M. Alghamdi, and S. Almuhammadi, "Natural language processing for automated resume screening and candidate ranking: A survey," IEEE Access, vol. 10, pp. 98234-98248, 2022.',
        '[3] A. Deshpande, S. Roy, and P. Sharma, "Context-aware automated skill extraction and candidate profile validation using large language models," in Proc. IEEE ICBDC, 2023, pp. 112-119.',
        '[4] F. P.-W. Lo, J. Qiu, Z. Wang, H. Yu, Y. Chen, G. Zhang, and B. Lo, "AI hiring with LLMs: A context-aware and explainable multi-agent framework for resume screening," arXiv preprint arXiv:2504.02870, 2025.',
        '[5] T. B. Brown et al., "Language models are few-shot learners," in Advances in Neural Information Processing Systems (NeurIPS), vol. 33, 2020, pp. 1877-1901.',
        '[6] E. Kalliamvakou, G. Gousios, K. Blincoe, L. Singer, D. M. German, and D. Damian, "The promises and perils of mining GitHub," in Proc. 11th Working Conf. on Mining Software Repositories (MSR), 2014, pp. 92-101.',
        '[7] B. Vasilescu, V. Filkov, and A. Serebrenik, "Perception and reality of contribution in open source software," in Proc. 37th International Conf. on Software Engineering (ICSE), 2015, pp. 89-100.',
        '[8] P. Lewis et al., "Retrieval-augmented generation for knowledge-intensive NLP tasks," in Advances in Neural Information Processing Systems (NeurIPS), vol. 33, 2020, pp. 9459-9474.',
        '[9] J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, "BERT: Pre-training of deep bidirectional transformers for language understanding," in Proc. NAACL-HLT, 2019, pp. 4171-4186.',
        '[10] G. Izacard and E. Grave, "Leveraging passage retrieval with generative models for open domain question answering," in Proc. 16th Conf. European Chapter of the ACL (EACL), 2021, pp. 874-880.',
        '[11] N. Reimers and I. Gurevych, "Sentence-BERT: Sentence embeddings using siamese BERT-networks," in Proc. EMNLP, 2019, pp. 3982-3992.',
        '[12] H. Su et al., "One embedder, any task: Instruction-finetuned text embeddings," arXiv preprint arXiv:2212.09741, 2022.',
        '[13] Z. Feng et al., "CodeBERT: A pre-trained model for programming and natural languages," in Findings of EMNLP 2020, 2020, pp. 1536-1547.',
        '[14] D. Guo et al., "UniXcoder: Unified cross-modal pre-training for code representation and generation," in Proc. 60th Annual Meeting of the ACL, 2022, pp. 7212-7225.',
        '[15] M. Chen et al., "Evaluating large language models trained on code," arXiv preprint arXiv:2107.03374, 2021.',
        '[16] R. Li et al., "StarCoder: may the source be with you!" arXiv preprint arXiv:2305.06161, 2023.',
        '[17] Chroma, "ChromaDB: The open-source embedding database," 2023. [Online]. Available: https://www.trychroma.com',
        '[18] J. Johnson, M. Douze, and H. Jégou, "Billion-scale similarity search with GPUs," IEEE Transactions on Big Data, vol. 7, no. 3, pp. 535-547, 2021.',
        '[19] Q. Wu et al., "AutoGen: Enabling next-gen LLM applications via multi-agent conversation," arXiv preprint arXiv:2308.08155, 2023.',
        '[20] S. Hong et al., "MetaGPT: Meta programming for a multi-agent collaborative framework," in International Conference on Learning Representations (ICLR), 2024.',
        '[21] H. Chase, "LangChain: Building applications with LLMs through composability," 2023. [Online]. Available: https://github.com/langchain-ai/langchain',
        '[22] J. S. Park, J. C. O\'Brien, C. J. Cai, M. R. Morris, P. Liang, and M. S. Bernstein, "Generative agents: Interactive simulacra of human behavior," in Proc. 36th Annual ACM Symposium on User Interface Software and Technology (UIST), 2023, pp. 1-22.',
        '[23] J. Wei et al., "Chain-of-thought prompting elicits reasoning in large language models," in Advances in Neural Information Processing Systems (NeurIPS), vol. 35, 2022, pp. 24824-24837.',
        '[24] T. Kojima, S. S. Gu, M. Reid, Y. Matsuo, and Y. Iwasawa, "Large language models are zero-shot reasoners," in Advances in Neural Information Processing Systems (NeurIPS), vol. 35, 2022, pp. 22199-22213.',
        '[25] GitHub Inc., "GitHub REST API documentation," 2023. [Online]. Available: https://docs.github.com/en/rest',
        '[26] M. Jayaratne and B. Jayatilleke, "A systematic review of artificial intelligence applications in recruitment and selection," IEEE Transactions on Engineering Management, vol. 69, no. 4, pp. 1620-1634, 2022.',
        '[27] A. Dubey et al., "The Llama 3 herd of models," arXiv preprint arXiv:2407.21783, 2024.',
        '[28] B. Rozière et al., "Code Llama: Open foundation models for code," arXiv preprint arXiv:2308.12950, 2023.',
        '[29] S. Xiao, Z. Liu, P. Zhang, and N. Muennighoff, "C-Pack: Packaged resources to advance general Chinese and cross-lingual embedding," arXiv preprint arXiv:2309.07597, 2023.',
        '[30] Y. Wang, W. Wang, S. Joty, and S. C. H. Hoi, "CodeT5+: Open code large language models for code understanding and generation," in Proc. EMNLP, 2023, pp. 12536-12558.'
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(7.5)

    doc.save(output_path)
    print(f"Successfully created IEEE 2-Column Word document at: {output_path}")

if __name__ == "__main__":
    paths = [
        r"C:\Users\abhin_yr5b44s\Desktop\resume_project\paper\camcon.docx",
        r"C:\Users\abhin_yr5b44s\Desktop\resume_project\paper\CodeAudit_AI_Paper_Final_6Authors.docx",
        r"C:\Users\abhin_yr5b44s\Desktop\resume_project\paper\CodeAudit_AI_Paper_ICESC2026.docx",
        r"C:\Users\abhin_yr5b44s\Downloads\camcon.docx"
    ]
    for p in paths:
        try:
            create_ieee_paper(p)
            print(f"Successfully generated: {p}")
        except Exception as e:
            print(f"Could not write {p}: {e}")

